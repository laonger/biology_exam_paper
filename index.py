#!/usr/bin/python
# encoding: utf-8


import re
import time
import random
import base64
import cStringIO
from PIL import Image
import json

from flask import Flask
from flask import request
from flask import send_file

from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE
from docx.enum.style import WD_STYLE_TYPE


import model

app = Flask(__name__)
app.debug = True


@app.route("/")
def index():
    """# index: docstring """
    return 'hi'
    

@app.route("/xuanzeti")
def xuanzeti():
    """# xuanzeti: 选择题列表 """
    l = []
    for i in model.Xuanzeti.list_all():
        l.append({
            'id'        : i[0],
            'timu'      : i[1],
            'peitu'     : i[2],
            'daan'      : i[3],
            'xueqi'     : i[4],
            'nandu'     : i[5],
        })
    
    return json.dumps({'xuanzeti': l})

@app.route("/xuanzeti_submit", methods=["POST"])
def xuanzeti_submit():
    """# xuanzeti: docstring """
    print request.json
    n = model.Xuanzeti.new(
        request.json.get('timu', ''),
        request.json.get('peitu', ''),
        request.json.get('daan', ''),
        request.json['xueqi'],
        int(request.json['nandu']),
    )
    return json.dumps({'k': 'OK'})

@app.route("/xuanzeti_delete/<id>", methods=["GET"])
def xuanzeti_delete(id):
    """# xuanzeti: docstring """
    model.Xuanzeti.delete(id)
    return "OK"
    
@app.route("/tiankongti")
def tiankongti():
    """# tiankongti: 选择题列表 """
    l = []
    for i in model.Tiankongti.list_all():
        l.append({
            'id'     : i[0],
            'timu'   : i[1],
            'peitu'  : i[2],
            'daan'   : i[3],
            'xueqi'  : i[4],
            'nandu'  : i[5],
        })
    
    return json.dumps({'tiankongti': l})

@app.route("/tiankongti_submit", methods=["POST"])
def tiankongti_submit():
    """# tiankongti: docstring """
    n = model.Tiankongti.new(
        request.json.get('timu', ''),
        request.json.get('peitu', ''),
        request.json.get('daan', ''),
        request.json['xueqi'],
        int(request.json['nandu']),
    )
    return json.dumps({'k': 'OK'})

@app.route("/tiankongti_delete/<id>", methods=["GET"])
def tiankongti_delete(id):
    """# tiankongti: docstring """
    model.Tiankongti.delete(id)
    return "OK"


@app.route("/chuti", methods=["POST"])
def chuti():
    """# chuti: docstring """
    data = request.form.to_dict(flat=False)
    print data
    xueqi_list = data.get('xueqi')
    if not xueqi_list:
        xueqi_list = []
        for i in 'ABCD':
            for j in xrange(1, 11):
                xueqi_list.append(i+str(j))

    xuanze_shuliang = int(data['xuanzeShuliang'][0])
    # 每个难度的比例, 下标是难度
    xuanze_nandu_list = [int(i) for i in data['xuanzeNandu']][::-1]

    tiankong_shuliang = int(data['tiankongShuliang'][0])
    # 每个难度的比例, 下标是难度
    tiankong_nandu_list = [int(i) for i in data['tiankongNandu']][::-1]

    xuanzeti_list = []
    tiankongti_list = []
    for ti_list, shuliang, nandu_list, model_class in [
        (xuanzeti_list, xuanze_shuliang, xuanze_nandu_list, model.Xuanzeti),
        (tiankongti_list, tiankong_shuliang, tiankong_nandu_list, model.Tiankongti),
    ]:

        ti_all = list(model_class.in_('xueqi', xueqi_list))
        if not ti_all:
            raise
        nandu_ratio_sum = sum(nandu_list)
        nandu_min = 10
        for nandu, ratio in enumerate(nandu_list):
            if not ratio:
                continue
            if nandu < nandu_min: nandu_min = nandu
            nandu_num = shuliang *ratio /nandu_ratio_sum
            temp = [i for i in ti_all if i.nandu == nandu]
            if len(temp) <= nandu_num:
                ti_list.extend(temp)
            else:
                ti_list.extend(random.sample(temp, nandu_num))
        if len(ti_list) < shuliang:
            nandu_num = shuliang - len(ti_list)
            temp = []
            ti_all.sort(key=lambda x: x.nandu)
            for i in ti_all:
                if i not in ti_list:
                    temp.append(i)
                if len(temp) >= nandu_num*2: break
            if len(temp) <= nandu_num:
                ti_list.extend(temp)
            else:
                ti_list.extend(random.sample(temp, nandu_num))

    
    print len(tiankongti_list), len(xuanzeti_list)

    d = Document()
    d.add_heading(unicode('标题', 'utf-8'), level=0)
    d.add_paragraph(' ')
    d.add_heading(unicode("一、选择题", 'utf-8'), level=1)
    d.add_paragraph(' ')
    xuanze_daan_list = []   # 题号，答案, None
    for n, i in enumerate(xuanzeti_list):
        p = d.add_paragraph(("%d. "%(n+1))+i.timu)
        xuanze_daan_list.append((n+1, i.daan, None))
        if i.peitu:
            pic = cStringIO.StringIO()
            pic_buffer = cStringIO.StringIO(base64.b64decode(i.peitu.split(',')[1]))
            image = Image.open(pic_buffer)
            image.save(pic, image.format, quality=100)
            pic.seek(0)
            d.add_picture(pic, width=Inches(3.0))
    d.add_paragraph(' ')
    d.add_heading(unicode("二、非选题", 'utf-8'), level=1)
    d.add_paragraph(' ')
    tiankong_daan_list = []     # 题号，答案，答题纸上的空白
    for n, i in enumerate(tiankongti_list):
        p = d.add_paragraph(("%d. "%(n+1))+i.timu)
        tihao = 0
        daan_length = 0

        # 准备答题纸上答题用的横线： (1)_____, (2)_____
        datizhi = []
        for j in re.findall(u'(\([1-9]+\))[^_]*(_+)', i.timu):
            datizhi.append(''.join(j))

        tiankong_daan_list.append((n+1, i.daan, ',   '.join(datizhi)))
        if i.peitu:
            pic = cStringIO.StringIO()
            pic_buffer = cStringIO.StringIO(base64.b64decode(i.peitu.split(',')[1]))
            image = Image.open(pic_buffer)
            image.save(pic, image.format, quality=100)
            pic.seek(0)
            d.add_picture(pic, width=Inches(3.0))

    # 答题纸
    d.add_page_break()    
    d.add_heading(unicode('答题纸', 'utf-8'), level=0)
    d.add_paragraph(' ')

    p = d.add_heading(unicode("姓名：                                            班级：          ", 'utf-8'), level=2)
    d.add_paragraph(' ')

    d.add_heading(unicode("一、选择题", 'utf-8'), level=1)
    d.add_paragraph(' ')
    cols = 10
    table = d.add_table(rows=2, cols=cols)
    table.style = 'Table Grid'
    i = 0
    cells_tihao = table.rows[0].cells
    cells_kongbai = table.rows[1].cells
    for v in xuanze_daan_list:
        cells_tihao[i].text = str(v[0])
        cells_kongbai[i].text = ' '
        i+=1
        if i == 10:
            i = 0
            cells_tihao = table.add_row().cells
            cells_kongbai = table.add_row().cells

    d.add_paragraph(' ')
    d.add_heading(unicode("二、非选题", 'utf-8'), level=1)
    d.add_paragraph(' ')
    for i in tiankong_daan_list:
        d.add_paragraph('%d. %s'%(i[0], i[2]))
    # 答案
    d.add_page_break()    
    d.add_heading(unicode('答案', 'utf-8'), level=0)
    d.add_paragraph(' ')
    d.add_heading(unicode("一、选择题", 'utf-8'), level=1)
    d.add_paragraph(' ')
    cols = 10
    table = d.add_table(rows=2, cols=cols)
    table.style = 'Table Grid'
    i = 0
    cells_tihao = table.rows[0].cells
    cells_kongbai = table.rows[1].cells
    for v in xuanze_daan_list:
        if i == 10:
            i = 0
            cells_tihao = table.add_row().cells
            cells_kongbai = table.add_row().cells
        cells_tihao[i].text = str(v[0])
        cells_kongbai[i].text = v[1]
        i+=1
    d.add_paragraph(' ')
    d.add_heading(unicode("二、非选题", 'utf-8'), level=1)
    d.add_paragraph(' ')
    for i in tiankong_daan_list:
        d.add_paragraph('%d. %s'%(i[0], i[1]))


    doc = cStringIO.StringIO()
    d.save(doc)
    doc.seek(0)
    
    return send_file(
        doc, 
        mimetype='appliaction/vnd.openxmlformats-officedocument.wordprocessingml.document',
        attachment_filename='shengwu_'+str(int(time.time()))+'.docx',
        as_attachment=True
    )



if __name__ == '__main__':
    app.run('0.0.0.0')
